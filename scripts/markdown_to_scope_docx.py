#!/usr/bin/env python3

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Mm(22)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(20)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(51, 65, 85)
    normal.paragraph_format.space_after = Pt(6)

    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(15, 23, 42)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(15, 23, 42)),
        ("Heading 2", 13, RGBColor(29, 78, 216)),
        ("Heading 3", 11.5, RGBColor(37, 99, 235)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_field_run(paragraph, instr_text: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    run._r.append(instr)

    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char)

    run = paragraph.add_run("1")
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def add_footer_with_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)
    add_field_run(paragraph, "PAGE")


def add_cover(document: Document) -> None:
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Documento de Alcance de Solucion")

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = meta.add_run("Solucion documental asistida por IA para implementacion con Mendix + Codex")
    r.font.name = "Aptos"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(71, 85, 105)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)

    info = [
        "Tipo de documento: Alcance funcional y tecnico",
        "Audiencia: Equipo de desarrollo, arquitectura y preventa",
        "Formato: Base para backlog, diseno funcional y definicion de implementacion",
    ]
    for line in info:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(line)

    document.add_page_break()


def add_toc_placeholder(document: Document) -> None:
    p = document.add_paragraph("Tabla de contenido", style="Heading 1")
    p.paragraph_format.space_after = Pt(8)
    toc = document.add_paragraph()
    add_field_run(toc, r'TOC \o "1-3" \h \z \u')
    note = document.add_paragraph()
    note_run = note.add_run("Actualiza la tabla de contenido en Word con clic derecho > Actualizar campo.")
    note_run.font.name = "Aptos"
    note_run.font.size = Pt(9)
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(100, 116, 139)
    document.add_page_break()


def add_inline_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.font.name = "Aptos"
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59)
        else:
            run = paragraph.add_run(token[2:-2])
            run.font.name = "Aptos"
            run.bold = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = "Aptos"


def body_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline_runs(p, text.strip())


def parse_markdown_to_docx(markdown: str, document: Document) -> None:
    lines = markdown.splitlines()
    paragraph_buffer: list[str] = []

    def flush_paragraph():
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
        if text:
            body_paragraph(document, text)
        paragraph_buffer.clear()

    for raw_line in lines:
        stripped = raw_line.strip()

        if not stripped:
            flush_paragraph()
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            p = document.add_paragraph(style="Heading 1")
            p.add_run(stripped[2:].strip())
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            p = document.add_paragraph(style="Heading 1")
            p.add_run(stripped[3:].strip())
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            p = document.add_paragraph(style="Heading 2")
            p.add_run(stripped[4:].strip())
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Mm(6)
            p.paragraph_format.first_line_indent = Mm(-3)
            add_inline_runs(p, stripped[2:].strip())
            continue

        ordered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if ordered:
            flush_paragraph()
            p = document.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Mm(6)
            p.paragraph_format.first_line_indent = Mm(-3)
            add_inline_runs(p, ordered.group(2).strip())
            continue

        if re.match(r"^[A-Za-z].*:\s*$", stripped) and len(stripped) < 70:
            flush_paragraph()
            p = document.add_paragraph(style="Normal")
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped)
            run.bold = True
            run.font.name = "Aptos"
            run.font.color.rgb = RGBColor(15, 23, 42)
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph()


def main() -> None:
    if len(sys.argv) != 3:
        print("Uso: markdown_to_scope_docx.py <input.md> <output.docx>", file=sys.stderr)
        raise SystemExit(1)

    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    target.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    set_document_defaults(document)
    add_footer_with_page_number(document.sections[0])
    add_cover(document)
    add_toc_placeholder(document)
    parse_markdown_to_docx(source.read_text(encoding="utf-8"), document)
    document.save(str(target))
    print(target)


if __name__ == "__main__":
    main()
